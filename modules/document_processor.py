import os
import re
import io
import logging
from typing import Optional

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text from PDF or DOCX files.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text as a string or None if extraction fails
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_ext}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}", exc_info=True)
        return None

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF files using PyPDF2.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        # Clean the extracted text
        return clean_text(text)
    except ImportError:
        # Fallback to pdfminer.six if PyPDF2 is not available
        return extract_text_from_pdf_with_pdfminer(file_path)
    except Exception as e:
        logger.error(f"Error extracting text from PDF with PyPDF2: {str(e)}", exc_info=True)
        # Try alternative method if primary method fails
        return extract_text_from_pdf_with_pdfminer(file_path)

def extract_text_from_pdf_with_pdfminer(file_path: str) -> str:
    """
    Extract text from PDF files using pdfminer.six.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        
        text = pdfminer_extract_text(file_path)
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error extracting text with pdfminer: {str(e)}", exc_info=True)
        raise

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
    """
    try:
        import docx
        
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return clean_text('\n'.join(full_text))
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
        raise

def clean_text(text: str) -> str:
    """
    Clean and normalize the extracted text.
    
    Args:
        text: The raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple newlines with single newline
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove non-breaking spaces and other special whitespace
    text = text.replace('\xa0', ' ')
    
    # Remove excessive spaces
    text = re.sub(r' +', ' ', text)
    
    # Remove page numbers and headers/footers (simplified approach)
    text = re.sub(r'\n\d+\n', '\n', text)
    
    # Remove common PDF artifacts
    text = re.sub(r'Form\s+\d+', '', text)
    
    return text.strip()
